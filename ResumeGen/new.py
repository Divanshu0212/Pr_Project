import os
import json
import base64
from fastapi import FastAPI, HTTPException, BackgroundTasks, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import google.generativeai as genai
from fpdf import FPDF, XPos, YPos
import tempfile
import requests
from io import BytesIO
from docx import Document
import re
from concurrent.futures import ThreadPoolExecutor
import logging # Import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # In production, replace with your frontend domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure Gemini API
# *** IMPORTANT: Replace "YOUR_GEMINI_API_KEY" with your actual key in a real environment ***
# It's best practice to use environment variables as shown
GOOGLE_API_KEY = "AIzaSyBjCoY3uemTDJdOTfFSWTb1lGoosUsvh_Y" # Replace placeholder if needed
if GOOGLE_API_KEY == "YOUR_GEMINI_API_KEY":
    logger.warning("Using placeholder Gemini API Key. Please set the GOOGLE_API_KEY environment variable.")
try:
    genai.configure(api_key=GOOGLE_API_KEY)
except Exception as e:
    logger.error(f"Failed to configure Gemini API: {e}")
    # You might want to prevent the app from starting fully if the API key is invalid
    # raise RuntimeError("Gemini API key configuration failed.") from e

# Models (Pydantic models remain the same)
class Education(BaseModel):
    institution: str
    degree: str
    field: str
    date_range: str
    gpa: Optional[str] = None

class Experience(BaseModel):
    company: str
    position: str
    date_range: str
    description: List[str]

class Project(BaseModel):
    name: str
    description: List[str]
    link: Optional[str] = None
    technologies: Optional[List[str]] = None

class Skill(BaseModel):
    category: str
    skills: List[str]

class Achievement(BaseModel):
    title: str
    description: str

class CodingProfile(BaseModel):
    platform: str
    url: str

class UserResumeData(BaseModel):
    name: str
    email: str
    phone: str
    linkedin: Optional[str] = None
    github: Optional[str] = None
    portfolio: Optional[str] = None
    target_profession: str
    education: List[Education]
    experiences: List[Experience] # Changed from Optional to required List
    projects: List[Project]
    skills: List[Skill]
    achievements: Optional[List[Achievement]] = None
    coding_profiles: Optional[List[CodingProfile]] = None

class OptimizationResponse(BaseModel):
    pdf_url: str
    ats_score: float
    optimized_data: Dict[str, Any]
    improvement_notes: List[str]

# Default resume data (remains the same)
default_resume_data = {
    "name": "Divanshu Bhargava",
    "email": "divanshubhargava026@gmail.com",
    "phone": "+91-9359992426",
    "linkedin": "Divanshu Bhargava",
    "github": "github.com/divanshu0212",
    "target_profession": "Software Engineer",
    "education": [
        {
            "institution": "Indian Institute Of Information Technology Jabalpur, India",
            "degree": "Bachelor of Technology",
            "field": "Computer Science and Engineering(CSE)",
            "date_range": "August 2023 - May 2027"
        },
        {
            "institution": "KRSD Public School Mathura, India",
            "degree": "CBSE Class XII",
            "field": "",
            "date_range": "April 2021 - June 2022",
            "gpa": "94.2"
        },
        {
            "institution": "Rajiv International School Mathura, India",
            "degree": "CBSE Class X",
            "field": "",
            "date_range": "April 2019 - May 2020",
            "gpa": "93.6"
        }
    ],
    "experiences": [], # Defaulting to empty list
    "projects": [
        {
            "name": "The Big Dawgs",
            "description": [
                "Developed a full-stack gaming community platform in 24 hours during a hackathon, integrating 20+ categorized sections and connecting to an API with 500,000+ games to enhance user engagement.",
                "Implemented an AI-powered game recommendation system using TfidfVectorizer and cosine similarity, achieving 85% accuracy in suggesting similar games based on user preferences and interactions.",
                "Optimized UI performance by 40% through responsive design implementation using Tailwind CSS and seamless adaptation across 15+ device sizes and orientations."
            ],
            "link": "https://github.com/Divanshu0212/BigDawgs",
            "technologies": ["React", "TfidfVectorizer", "Tailwind CSS"]
        },
        {
            "name": "BigDocs",
            "description": [
                "Architected and deployed a comprehensive healthcare platform in 48 hours during a hackathon, implementing 7+ core features including telemedicine, appointment management, and AI-powered disease prediction.",
                "Developed and integrated a custom BERT-based disease prediction model achieving 92% accuracy across 100+ common diseases, processing symptom inputs in under 3 seconds for rapid diagnosis assistance.",
                "Engineered a real-time appointment system handling 50+ concurrent users, reducing appointment scheduling time by 75% and improving doctor-patient communication efficiency by 60%.",
                "Implemented a secure telemedicine infrastructure supporting 1080p video quality with 99.9% uptime, integrated with a prescription management system capable of handling 1000+ medical records daily."
            ],
            "link": "https://github.com/Divanshu0212/BigDocs",
            "technologies": ["BERT", "Telemedicine", "Healthcare"]
        },
        {
            "name": "Diggy",
            "description": [
                "Engineered a full-stack food delivery platform using MERN stack and Vite, reducing initial load time by 65% compared to traditional Create-React-App, handling 100+ concurrent user sessions.",
                "Designed and implemented 25+ reusable React components with Tailwind CSS, achieving 99% responsive design across 15+ device sizes and decreasing development time by 40%.",
                "Developed a RESTful API using Node.js and Express.js, enabling seamless communication between 50+ restaurants and 1,000+ food items.",
                "Optimized API performance by implementing caching, reducing response time by 70% and database load by 45%."
            ],
            "link": "https://github.com/Divanshu0212/Diggy",
            "technologies": ["MERN Stack", "Vite", "React", "Tailwind CSS", "Node.js", "Express.js"]
        }
    ],
    "skills": [
        {
            "category": "Programming Languages",
            "skills": ["Python", "PHP", "C++", "JavaScript", "TypeScript", "SQL", "Java"]
        },
        {
            "category": "Frameworks",
            "skills": ["React", "Next.js", "Redux", "Node.js", "Express.js", "Scikit-learn", "TensorFlow", "Keras"]
        },
        {
            "category": "Tools",
            "skills": ["Git", "Docker", "JWT", "Nodemon", "Web3.js", "MySQL", "MongoDB", "NumPy", "Pandas", "Matplotlib", "Plotly"]
        },
        {
            "category": "Platforms",
            "skills": ["Linux", "Windows", "AWS", "GCP"]
        },
        {
            "category": "Soft Skills",
            "skills": ["Leadership", "Event Management", "Writing", "Public Speaking", "Time Management"]
        }
    ],
    "achievements": [
        {
            "title": "Coding Achievement",
            "description": "Engineered solutions for over 375 data structure and algorithms problems on coding platforms, improving rating by 600 points and refining skills in time complexity analysis."
        },
        {
            "title": "Practice Consistency",
            "description": "Resolved over 200 competitive programming challenges on CodeChef, 100 on Codeforces, and 75 on LeetCode, showcasing a commitment to improving algorithmic expertise."
        },
        {
            "title": "Vaad-Vivaad",
            "description": "Secured 1st place in a city-wide debate competition, outshining 20+ competing teams with exceptional argumentation, eloquence, and critical thinking."
        },
        {
            "title": "JEE Main",
            "description": "Ranked in the top 1% in the JEE Main exam, with a percentile of 98.83 and an AIR rank of 13679."
        },
        {
            "title": "Robo Rush 2024",
            "description": "Secured 5th position in Robo Rush 2024 conducted by ERS club of IIITDMJ among 50+ teams, where we built a line-following and remote-controlled robot car."
        }
    ],
    "coding_profiles": [
        {
            "platform": "CodeChef",
            "url": "https://www.codechef.com/users/divanshu0212"
        },
        {
            "platform": "Codeforces",
            "url": "https://codeforces.com/profile/divanshu0212"
        },
        {
            "platform": "Leetcode",
            "url": "https://leetcode.com/u/divanshu0212/"
        }
    ]
}

# Helper functions
def get_gemini_model():
    """Safely get the Gemini model instance."""
    try:
        # Check if API key is configured before creating model
        if not GOOGLE_API_KEY or GOOGLE_API_KEY == "GEMINI_API_KEY":
            logger.error("Gemini API key is not configured properly.")
            return None
        return genai.GenerativeModel('gemini-1.5-flash')
    except Exception as e:
        logger.error(f"Error creating Gemini model: {e}")
        return None

def extract_json_from_text(text):
    """Attempts to extract a JSON object or array from a string,
       handling markdown code blocks and potential surrounding text."""
    # Try finding JSON within markdown code blocks first
    match = re.search(r'```(?:json)?\s*([\[\{].*?[\]\}])\s*```', text, re.DOTALL)
    if match:
        json_str = match.group(1)
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.warning(f"Found markdown JSON but failed to parse: {e}. Content: {json_str[:100]}...")
            # Fall through to try finding JSON anywhere

    # If no markdown JSON or parsing failed, try finding the first '{' or '['
    # and the last '}' or ']'
    start_index = -1
    end_index = -1
    brace_index = text.find('{')
    bracket_index = text.find('[')

    if brace_index != -1 and bracket_index != -1:
        start_index = min(brace_index, bracket_index)
    elif brace_index != -1:
        start_index = brace_index
    elif bracket_index != -1:
        start_index = bracket_index

    if start_index != -1:
        # Find the corresponding closing bracket/brace
        # This simple find might not be perfect for nested structures,
        # but it's better than nothing if markdown fails.
        # Consider using a stack-based approach for more complex cases if needed.
        end_brace = text.rfind('}')
        end_bracket = text.rfind(']')
        end_index = max(end_brace, end_bracket)

        if end_index > start_index:
            json_str = text[start_index : end_index + 1]
            try:
                return json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.warning(f"Found potential JSON structure but failed to parse: {e}. Content: {json_str[:100]}...")

    # If nothing worked
    logger.warning("Could not extract valid JSON from text.")
    return None


def optimize_content_with_gemini(section_name, content, target_profession):
    """Optimize content using Gemini API for better ATS scores with robust JSON parsing"""
    model = get_gemini_model()
    if not model:
         logger.error(f"Skipping optimization for {section_name} due to model initialization failure.")
         return content # Return original content if model failed

    # Craft specialized prompts based on section type
    if section_name == "projects":
        prompt = f"""
        You are an expert ATS (Applicant Tracking System) optimizer for resumes.
        Optimize these projects for a "{target_profession}" position with special focus on expanding descriptions.

        Original projects (JSON format):
        {json.dumps(content, indent=2)}

        Instructions:
        1. For each project, expand the description with more technical details and measurable outcomes.
        2. Add 1-2 more bullet points per project if possible, focusing on:
           - Specific technologies used and how they were applied
           - Challenges faced and how they were overcome
           - Quantifiable results (performance improvements, user growth, etc.)
           - Team size and your specific contributions
        3. Maintain all original factual information - only add details that could reasonably be inferred.
        4. Use strong action verbs relevant to "{target_profession}".
        5. Integrate relevant keywords naturally.
        6. Keep the same JSON structure - only return the optimized projects array.

        Optimized projects (JSON only):
        """
    elif section_name == "achievements":
        prompt = f"""
        You are an expert ATS (Applicant Tracking System) optimizer for resumes.
        Enhance these achievements for a "{target_profession}" position by adding more context and impact.

        Original achievements (JSON format):
        {json.dumps(content, indent=2)}

        Instructions:
        1. For each achievement, expand the description to include:
           - The context or competition level (e.g., "out of 500 participants")
           - Specific skills demonstrated
           - Quantifiable impact or results
           - Any recognition or awards received
        2. Maintain all original information while adding relevant details.
        3. Use strong action verbs and professional language.
        4. Keep the same JSON structure - only return the optimized achievements array.

        Optimized achievements (JSON only):
        """
    

    try:
        logger.info(f"Optimizing section: {section_name} for profession: {target_profession}")
        response = model.generate_content(prompt)
        response_text = response.text

        logger.debug(f"Raw Gemini response for {section_name}:\n{response_text}")

        optimized_content = extract_json_from_text(response_text)

        if optimized_content:
            # Basic validation
            if isinstance(content, list) and not isinstance(optimized_content, list):
                 logger.warning(f"JSON structure mismatch for {section_name}. Expected list, got {type(optimized_content)}. Returning original.")
                 return content
            if isinstance(content, dict) and not isinstance(optimized_content, dict):
                 logger.warning(f"JSON structure mismatch for {section_name}. Expected dict, got {type(optimized_content)}. Returning original.")
                 return content

            logger.info(f"Successfully parsed optimized JSON for {section_name}")
            return optimized_content
        else:
            logger.error(f"Failed to extract valid JSON from response for {section_name}. Returning original content.")
            return content

    except Exception as e:
        logger.error(f"Error during Gemini API call for {section_name}: {e}", exc_info=True)
        return content
    

# calculate_ats_score function (using robust parsing)
def calculate_ats_score(resume_data, target_profession):
    """Calculate an estimated ATS score based on content with focus on detailed descriptions"""
    model = get_gemini_model()
    if not model:
        logger.error("Cannot calculate ATS score, Gemini model failed to initialize.")
        return 50.0

    # Create a version that highlights description quality
    evaluation_data = {
        "target_profession": resume_data.get("target_profession", "N/A"),
        "skills_match": [f"{s['category']}: {len(s['skills'])} skills" for s in resume_data.get("skills", [])],
        "project_description_quality": [
            f"Project '{p['name']}': {len(p.get('description', []))} points, avg length {sum(len(d) for d in p.get('description', []))/max(1, len(p.get('description', [])))} chars"
            for p in resume_data.get("projects", [])
        ],
        "experience_description_quality": [
            f"Experience at {e['company']}': {len(e.get('description', []))} points"
            for e in resume_data.get("experiences", [])
        ],
        "achievement_detail_level": len(resume_data.get("achievements", []))
    }

    prompt = f"""
    You are an expert ATS evaluator. Score this resume (1-100) for "{target_profession}".
    Pay special attention to:
    - Depth and detail in project/experience descriptions
    - Quantification of achievements
    - Specificity of skills

    Evaluation Data:
    {json.dumps(evaluation_data, indent=2)}

    Scoring Guidelines:
    - 90-100: Exceptional detail, strong metrics, perfect skill match
    - 80-89: Strong descriptions, good metrics, good skill match
    - 70-79: Adequate descriptions, some metrics, basic skill match
    - Below 70: Needs improvement in descriptions or relevance

    Return ONLY the score (0-100) as a number. No other text.
    """

    try:
        logger.info(f"Calculating detailed ATS score for {target_profession}")
        response = model.generate_content(prompt)
        score_text = response.text.strip()
        
        # Extract score
        score_match = re.search(r'(\d{1,3})', score_text)
        if score_match:
            score = float(score_match.group(1))
            return max(0, min(100, score))
        return 75.0  # Default if parsing fails
    except Exception as e:
        logger.error(f"Error calculating detailed ATS score: {e}")
        return 75.0
    
    
# generate_improvement_notes function (using robust parsing)
def generate_improvement_notes(original_data, optimized_data, target_profession):
    """Generate notes on what was improved"""
    model = get_gemini_model()
    if not model:
         logger.error("Cannot generate improvement notes, Gemini model failed to initialize.")
         return ["Could not connect to AI to generate notes."]

    # Create summaries for comparison to avoid overly long prompts
    def summarize_section(data, section_key):
        if section_key not in data or not data[section_key]:
            return f"No {section_key} provided."
        if section_key in ['experiences', 'projects']:
            return f"{len(data[section_key])} {section_key} entries. First entry summary: {json.dumps(data[section_key][0].get('description', ['N/A'])[0][:100])}..."
        if section_key == 'skills':
             return ", ".join([f"{s['category']}: {len(s['skills'])} skills" for s in data[section_key]])
        # Add summaries for other sections if needed
        return f"{len(data[section_key])} entries in {section_key}."

    original_summary = {k: summarize_section(original_data, k) for k in ['experiences', 'projects', 'skills', 'achievements'] if k in original_data}
    optimized_summary = {k: summarize_section(optimized_data, k) for k in ['experiences', 'projects', 'skills', 'achievements'] if k in optimized_data}


    prompt = f"""
    You are an expert resume consultant comparing an original resume summary to an AI-optimized version for a "{target_profession}" position.
    Based SOLELY on the differences implied by these summaries, provide 3-5 concise bullet points (as a JSON array of strings) explaining the KEY improvements likely made by the AI for ATS compatibility.

    Focus on common ATS optimization techniques:
    - Stronger action verbs.
    - Quantification of achievements.
    - Keyword integration for "{target_profession}".
    - Clarity and conciseness.

    Original Summary: {json.dumps(original_summary, indent=2)}

    Optimized Summary: {json.dumps(optimized_summary, indent=2)}

    **Return ONLY a JSON array of strings, where each string is a brief improvement note (1-2 sentences). Example: ["Note 1", "Note 2", ...]**
    Notes:
    """

    default_notes = [
        "Optimized action verbs for stronger impact.",
        "Enhanced technical keyword density relevant to the target profession.",
        "Quantified achievements more effectively where possible.",
        "Improved overall phrasing for ATS readability."
    ]

    try:
        logger.info(f"Generating improvement notes for profession: {target_profession}")
        response = model.generate_content(prompt)
        notes_text = response.text.strip()
        logger.debug(f"Raw improvement notes response: {notes_text}")

        # Use the improved extraction function
        notes = extract_json_from_text(notes_text)

        if notes and isinstance(notes, list) and all(isinstance(n, str) for n in notes):
            logger.info("Successfully parsed improvement notes.")
            return notes
        else:
             logger.error(f"Failed to extract valid JSON array of strings for improvement notes. Response: '{notes_text}'. Returning default notes.")
             return default_notes

    except Exception as e:
        logger.error(f"Error generating improvement notes: {e}", exc_info=True)
        return default_notes # Default notes on error

# create_resume_pdf function (remains the same)
def create_resume_pdf(resume_data):
    """Create a professionally formatted resume PDF using only built-in fonts"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Use built-in helvetica font and ASCII bullets
    current_font = "helvetica"
    bullet = "- "  # ASCII bullet
    
    # Set initial font
    pdf.set_font(current_font, "", 16)
    
    # Name (centered)
    pdf.cell(0, 10, resume_data["name"], new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    
    # Contact info
    pdf.set_font(current_font, "", 9)
    contact_parts = [
        resume_data.get('email'),
        resume_data.get('phone'),
        f"LinkedIn: {resume_data['linkedin']}" if resume_data.get('linkedin') else None,
        f"GitHub: {resume_data['github']}" if resume_data.get('github') else None,
        f"Portfolio: {resume_data['portfolio']}" if resume_data.get('portfolio') else None
    ]
    contact_info = " | ".join(filter(None, contact_parts))
    pdf.cell(0, 6, contact_info, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(5)

    # Function to add a section title
    def add_section_title(title):
        pdf.set_font(current_font, "B", 12)
        pdf.cell(0, 8, title.upper(), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_line_width(0.5)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)

    # --- Education ---
    if resume_data.get("education"):
        add_section_title("Education")
        pdf.set_font(current_font, "", 10)
        for edu in resume_data["education"]:
            # Institution and date
            pdf.set_font(current_font, "B", 10)
            pdf.cell(0, 6, edu.get("institution", "N/A"))
            pdf.set_font(current_font, "", 10)
            pdf.cell(0, 6, edu.get("date_range", "N/A"), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            # Degree and field
            degree_info = edu.get('degree', 'N/A')
            if edu.get('field'):
                degree_info += f" in {edu['field']}"
            pdf.cell(0, 6, degree_info, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            # GPA if available
            if edu.get("gpa"):
                pdf.set_font(current_font, "I", 9)
                pdf.cell(0, 5, f"GPA/Percentage: {edu['gpa']}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

    # --- Skills ---
    if resume_data.get("skills"):
        add_section_title("Skills")
        pdf.set_font(current_font, "", 10)
        for skill_group in resume_data["skills"]:
            pdf.set_font(current_font, "B", 10)
            pdf.cell(40, 6, f"{skill_group.get('category', 'N/A')}:")
            pdf.set_font(current_font, "", 10)
            skills_text = ", ".join(skill_group.get("skills", []))
            pdf.multi_cell(0, 6, skills_text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)

    # --- Experience ---
    if resume_data.get("experiences"):
        add_section_title("Experience")
        for exp in resume_data["experiences"]:
            # Position and date
            pdf.set_font(current_font, "B", 10)
            pdf.cell(0, 6, exp.get("position", "N/A"))
            pdf.set_font(current_font, "", 10)
            pdf.cell(0, 6, exp.get("date_range", "N/A"), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            # Company
            pdf.set_font(current_font, "I", 10)
            pdf.cell(0, 6, exp.get("company", "N/A"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            # Description bullets
            pdf.set_font(current_font, "", 10)
            for desc in exp.get("description", []):
                pdf.set_x(pdf.l_margin + 10)  # Indent bullet points
                pdf.multi_cell(0, 5, f"{bullet}{desc}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(0.5)
            pdf.ln(2)

    # --- Projects ---
    if resume_data.get("projects"):
        add_section_title("Projects")
        for project in resume_data["projects"]:
            # Project name and link
            pdf.set_font(current_font, "B", 10)
            pdf.cell(0, 6, project.get("name", "N/A"))
            if project.get("link"):
                pdf.set_font(current_font, "", 9)
                pdf.set_text_color(0, 0, 255)  # Blue for links
                pdf.cell(0, 6, project["link"], align="R", link=project["link"])
                pdf.set_text_color(0, 0, 0)  # Reset color
            pdf.set_font(current_font, "", 10)
            pdf.ln(6) if project.get("link") else pdf.ln(6)
            
            # Description bullets
            for desc in project.get("description", []):
                pdf.set_x(pdf.l_margin + 10)
                pdf.multi_cell(0, 5, f"{bullet}{desc}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(0.5)
            
            # Technologies if available
            if project.get("technologies"):
                pdf.set_x(pdf.l_margin + 10)
                pdf.set_font(current_font, "I", 9)
                tech_text = "Technologies: " + ", ".join(project['technologies'])
                pdf.multi_cell(0, 5, tech_text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

    # --- Achievements ---
    if resume_data.get("achievements"):
        add_section_title("Achievements")
        pdf.set_font(current_font, "", 10)
        for ach in resume_data["achievements"]:
            pdf.set_font(current_font, "B", 10)
            pdf.multi_cell(0, 5, f"{bullet}{ach.get('title', 'N/A')}:", 
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(current_font, "", 10)
            pdf.set_x(pdf.l_margin + 10)  # Indent description
            pdf.multi_cell(0, 5, ach.get("description", "N/A"), 
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)

    # --- Coding Profiles ---
    if resume_data.get("coding_profiles"):
        add_section_title("Coding Profiles")
        pdf.set_font(current_font, "", 10)
        for profile in resume_data["coding_profiles"]:
            pdf.set_x(pdf.l_margin + 10)
            pdf.cell(40, 5, f"{bullet}{profile.get('platform', 'N/A')}:")
            pdf.set_text_color(0, 0, 255)  # Blue for links
            pdf.cell(0, 5, profile.get("url", "#"), link=profile.get("url", "#"))
            pdf.set_text_color(0, 0, 0)  # Reset color
            pdf.ln(5)

    # Save to a temporary file
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            pdf_path = tmp.name
            pdf.output(pdf_path)
            logger.info(f"Generated PDF at: {pdf_path}")
        return pdf_path
    except Exception as e:
        logger.error(f"Failed to generate or save PDF: {e}", exc_info=True)
        return None

def optimize_resume_section(section_name, content, target_profession):
    """Wrapper function for parallel optimization of resume sections"""
    # Ensure content is not empty before calling API
    if not content:
         logger.info(f"Skipping optimization for empty section: {section_name}")
         return section_name, content
    return section_name, optimize_content_with_gemini(section_name, content, target_profession)

# API endpoints
@app.post("/api/optimize-resume")
async def optimize_resume(resume_data: UserResumeData, background_tasks: BackgroundTasks):
    """Optimize resume content for better ATS scores"""
    try:
        original_data = resume_data.model_dump()
        target_profession = original_data.get("target_profession", "Software Engineer") # Default if missing
        logger.info(f"Received optimization request for profession: {target_profession}")

        # Create a copy for optimization
        optimizing_data = original_data.copy()

        # Prepare sections for parallel optimization
        # Only include sections that actually have content
        sections_to_optimize_inputs = []
        if optimizing_data.get("experiences"):
            sections_to_optimize_inputs.append(("experiences", optimizing_data["experiences"]))
        if optimizing_data.get("projects"):
            sections_to_optimize_inputs.append(("projects", optimizing_data["projects"]))
        if optimizing_data.get("skills"):
             sections_to_optimize_inputs.append(("skills", optimizing_data["skills"]))
        if optimizing_data.get("achievements"):
             sections_to_optimize_inputs.append(("achievements", optimizing_data["achievements"]))

        logger.info(f"Optimizing sections: {[s[0] for s in sections_to_optimize_inputs]}")

        # Optimize sections in parallel
        if sections_to_optimize_inputs: # Only run thread pool if there's something to optimize
            with ThreadPoolExecutor(max_workers=4) as executor:
                # Submit tasks
                futures = [
                    executor.submit(optimize_resume_section, section_name, content, target_profession)
                    for section_name, content in sections_to_optimize_inputs
                ]

                # Collect results
                for future in futures:
                    try:
                        section_name, optimized_content = future.result()
                        optimizing_data[section_name] = optimized_content
                        logger.info(f"Completed optimization for section: {section_name}")
                    except Exception as exc:
                        # Log error from thread but don't necessarily stop the whole process
                        # The section will retain its original content due to error handling within optimize_resume_section
                        logger.error(f'Optimization thread for a section generated an exception: {exc}', exc_info=True)

        # Calculate ATS score for potentially partially or fully optimized resume
        ats_score = calculate_ats_score(optimizing_data, target_profession)

        # Generate improvement notes based on original vs (partially) optimized
        improvement_notes = generate_improvement_notes(original_data, optimizing_data, target_profession)
        print(improvement_notes)

        # Generate PDF from the (partially) optimized data
        pdf_path = create_resume_pdf(optimizing_data)
        if not pdf_path:
             raise HTTPException(status_code=500, detail="Failed to generate the resume PDF.")


        # Encode PDF as base64
        try:
            with open(pdf_path, "rb") as pdf_file:
                pdf_base64 = base64.b64encode(pdf_file.read()).decode("utf-8")
            pdf_url = f"data:application/pdf;base64,{pdf_base64}"
            logger.info("Successfully encoded PDF to base64.")
        except Exception as e:
            logger.error(f"Failed to read or encode PDF: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail="Failed to read or encode the generated PDF.")
        finally:
            # Clean up the temporary file in the background, even if encoding failed
             if pdf_path and os.path.exists(pdf_path):
                 background_tasks.add_task(os.unlink, pdf_path)
                 logger.info(f"Scheduled cleanup for temporary PDF: {pdf_path}")

        return OptimizationResponse(
            pdf_url=pdf_url,
            ats_score=ats_score,
            optimized_data=optimizing_data, # Send back the potentially partially optimized data
            improvement_notes=improvement_notes
        )

    except HTTPException as http_exc:
         logger.error(f"HTTP Exception during optimization: {http_exc.detail}", exc_info=True)
         raise http_exc # Re-raise FastAPI exceptions
    except Exception as e:
        logger.error(f"Unexpected error in optimize_resume endpoint: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")


@app.get("/api/default-resume")
async def get_default_resume():
    """Provide default resume data for the frontend"""
    logger.info("Serving default resume data.")
    # Ensure experiences is always a list in the default data served
    data_to_serve = default_resume_data.copy()
    if 'experiences' not in data_to_serve or data_to_serve['experiences'] is None:
        data_to_serve['experiences'] = []
    return data_to_serve


@app.post("/api/upload-resume")
async def upload_resume(
    file: UploadFile = File(...),
    target_profession: str = Form("Software Engineer") # Default target
):
    """Parse an uploaded resume file and extract information (placeholder implementation)"""
    logger.info(f"Received resume upload: {file.filename}, target profession: {target_profession}")
    content_type = file.content_type
    file_ext = os.path.splitext(file.filename)[-1].lower() if file.filename else ""

    # --- Placeholder Logic ---
    # In a real app, you'd use libraries like PyMuPDF for PDFs, python-docx for DOCX
    # and then potentially send the extracted text to Gemini for structuring.
    # For now, we just return the default data but set the target profession.
    logger.warning("Resume upload parsing is currently a placeholder. Returning default data.")

    parsed_data = default_resume_data.copy() # Start with default
    parsed_data["target_profession"] = target_profession # Set target profession

    # You *could* add basic text extraction here if needed for debugging
    # but full parsing is complex.
    # Example (needs python-docx installed: pip install python-docx):
    # if file_ext == ".docx":
    #    try:
    #        doc_content = await file.read()
    #        doc = Document(BytesIO(doc_content))
    #        full_text = "\n".join([para.text for para in doc.paragraphs])
    #        logger.info(f"Extracted text from DOCX (first 500 chars): {full_text[:500]}")
    #        # ---> Here you would call Gemini to parse 'full_text' into JSON <---
    #    except Exception as e:
    #        logger.error(f"Failed basic DOCX read for {file.filename}: {e}")


    return parsed_data
    # --- End Placeholder Logic ---

    # try:
    #     # Read uploaded file content
    #     content = await file.read()

    #     # Process based on file type
    #     parsed_data = default_resume_data.copy() # Start with default data
    #     parsed_data["target_profession"] = target_profession

    #     file_ext = os.path.splitext(file.filename)[-1].lower() if file.filename else ""
    #     logger.info(f"Processing file with extension: {file_ext}")

    #     # Add actual parsing logic here using appropriate libraries (PyMuPDF, python-docx)
    #     if file_ext == ".pdf":
    #         # Use PyMuPDF (fitz) to extract text
    #         # Send text to Gemini to structure into JSON matching UserResumeData
    #         logger.info("PDF parsing logic not implemented yet.")
    #         pass # Placeholder

    #     elif file_ext in [".doc", ".docx"]:
    #         # Use python-docx to extract text
    #         # Send text to Gemini to structure into JSON matching UserResumeData
    #         logger.info("DOC/DOCX parsing logic not implemented yet.")
    #         try:
    #             doc = Document(BytesIO(content))
    #             full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    #             logger.info(f"Extracted DOCX text (first 300 chars): {full_text[:300]}")
    #             # ---> Call Gemini here <---
    #         except Exception as docx_err:
    #              logger.error(f"Error parsing DOCX file: {docx_err}")
    #         pass # Placeholder

    #     else:
    #         logger.warning(f"Unsupported file type uploaded: {file_ext}")
    #         # Optionally raise an error or just return default
    #         # raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")

    #     # Ensure target profession is set
    #     parsed_data["target_profession"] = target_profession

    #     return parsed_data

    # except Exception as e:
    #     logger.error(f"Failed to parse uploaded resume: {str(e)}", exc_info=True)
    #     raise HTTPException(status_code=500, detail=f"Failed to parse resume: {str(e)}")

@app.post("/api/analyze-job-description")
async def analyze_job_description(job_description: str = Form(...), current_resume: str = Form(...)): # Receive resume as JSON string
    """Analyze a job description and provide tailoring suggestions"""
    logger.info("Received request to analyze job description.")
    model = get_gemini_model()
    if not model:
         raise HTTPException(status_code=503, detail="AI Model Service Unavailable")

    try:
        # Parse current resume JSON string
        try:
             current_resume_data = json.loads(current_resume)
        except json.JSONDecodeError:
             logger.error("Failed to parse current_resume JSON string in analyze endpoint.")
             raise HTTPException(status_code=400, detail="Invalid format for current_resume data.")

        # It's good practice to summarize the resume data for the prompt if it's large
        # (Similar to the summarize_section function used in generate_improvement_notes)
        # For simplicity here, we send it as is, but be mindful of prompt size limits.

        prompt = f"""
        Analyze the provided Job Description and the Current Resume data for a potential candidate.
        Your goal is to assess the match and provide actionable feedback for tailoring the resume.

        Job Description:
        ---
        {job_description}
        ---

        Current Resume Data (JSON):
        ---
        {json.dumps(current_resume_data, indent=2)}
        ---

        Instructions:
        Return a JSON object containing the following fields ONLY:
        1.  "match_score": An estimated percentage (0-100) indicating how well the *current* resume matches the job description requirements (keywords, skills, experience level).
        2.  "key_missing_skills": A JSON array of strings listing specific, important skills or qualifications mentioned in the job description that seem to be MISSING or underrepresented in the resume. Limit to the top 5 most critical.
        3.  "tailoring_suggestions": A JSON array of strings providing 3-5 concrete, actionable suggestions for how the candidate could tailor their *existing* resume content (e.g., experiences, projects) to better align with this specific job description. Focus on rephrasing, highlighting relevant aspects, and using keywords from the job description.
        4.  "keyword_emphasis": A JSON array of strings listing 3-5 keywords or phrases *already present* in the resume that are particularly relevant to this job description and should be emphasized or potentially elaborated upon.

        **CRITICAL: Respond ONLY with the valid JSON object described above. No extra text, explanations, or markdown.**

        Analysis Result (JSON only):
        """

        logger.info("Sending job description analysis request to Gemini.")
        response = model.generate_content(prompt)
        analysis_text = response.text.strip()
        logger.debug(f"Raw job analysis response: {analysis_text}")

        # Use the robust JSON extractor
        analysis = extract_json_from_text(analysis_text)

        if analysis and isinstance(analysis, dict) and all(k in analysis for k in ["match_score", "key_missing_skills", "tailoring_suggestions", "keyword_emphasis"]):
            logger.info("Successfully parsed job description analysis.")
            # Optional: Validate inner types (e.g., ensure lists contain strings)
            return analysis
        else:
            logger.error(f"Failed to extract valid JSON for job analysis. Response: {analysis_text}")
            # Return a structured error or default analysis
            return {
                 "match_score": 50,
                 "key_missing_skills": ["Analysis Error: Could not determine missing skills."],
                 "tailoring_suggestions": ["Analysis Error: Could not generate suggestions. Review JD manually."],
                 "keyword_emphasis": ["Analysis Error: Could not determine keywords."]
            }

    except HTTPException as http_exc:
        raise http_exc # Re-raise validation errors etc.
    except Exception as e:
        logger.error(f"Failed during job description analysis: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to analyze job description: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    # Ensure GOOGLE_API_KEY is set before running
    if GOOGLE_API_KEY == "YOUR_GEMINI_API_KEY":
         print("\n*** WARNING: Running with placeholder GOOGLE_API_KEY. Optimization features will likely fail. ***")
         print("*** Please set the GOOGLE_API_KEY environment variable or replace the placeholder in the script. ***\n")

    uvicorn.run(app, host="0.0.0.0", port=8000)